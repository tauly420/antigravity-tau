"""
DOCX rendering pipeline for lab reports.

Converts LaTeX math expressions to OMML via latex2mathml + MML2OMML.XSL,
handles Hebrew RTL text direction, and generates DOCX output via python-docx.
"""

import base64
import logging
import os
import re
from io import BytesIO
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

import latex2mathml.converter

logger = logging.getLogger(__name__)

# ---------- Module-level constants ----------

UTILS_DIR = os.path.dirname(os.path.abspath(__file__))
XSL_PATH = os.path.join(UTILS_DIR, 'MML2OMML.XSL')

# Load XSLT once at module level
_xslt_tree = etree.parse(XSL_PATH)
_xslt_transform = etree.XSLT(_xslt_tree)

MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
WORD_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

HEADERS = {
    'he': {
        'theory': '1. רקע תיאורטי',
        'method': '2. שיטת מדידה',
        'results': '3. תוצאות',
        'discussion': '4. דיון',
        'conclusions': '5. מסקנות',
    },
    'en': {
        'theory': '1. Theoretical Background',
        'method': '2. Measurement Method',
        'results': '3. Results',
        'discussion': '4. Discussion',
        'conclusions': '5. Conclusions',
    },
}

DEFAULT_FONT = 'Noto Sans Hebrew'
DEFAULT_FONT_SIZE = Pt(11)

# Inline math regex: $...$ but not $$
INLINE_MATH_RE = re.compile(r'(?<!\$)\$(?!\$)([^$\n]+?)\$(?!\$)')


# ---------- Core functions ----------

def latex_to_omml(latex_expr: str) -> etree._Element:
    """Convert LaTeX expression to OMML element.

    Uses latex2mathml to get MathML, then applies MML2OMML.XSL XSLT
    to produce OMML. For inline math, extracts <m:oMath> from any
    wrapping <m:oMathPara>.

    Args:
        latex_expr: Raw LaTeX string (without $ delimiters).

    Returns:
        lxml Element with tag {math_ns}oMath.

    Raises:
        Exception on conversion failure.
    """
    # Step 1: LaTeX -> MathML
    mathml_str = latex2mathml.converter.convert(latex_expr)

    # Step 2: MathML -> OMML via XSLT
    mathml_tree = etree.fromstring(mathml_str.encode('utf-8'))
    omml_tree = _xslt_transform(mathml_tree)

    root = omml_tree.getroot()

    # Extract oMath from oMathPara wrapper if present (inline math)
    omath_para_tag = f'{{{MATH_NS}}}oMathPara'
    omath_tag = f'{{{MATH_NS}}}oMath'

    if root.tag == omath_para_tag:
        omath = root.find(f'{{{MATH_NS}}}oMath')
        if omath is not None:
            return deepcopy(omath)
    if root.tag == omath_tag:
        return deepcopy(root)

    # Search children
    for child in root.iter():
        if child.tag == omath_tag:
            return deepcopy(child)

    # If we got here but have valid XML, wrap in oMath
    omath = etree.Element(f'{{{MATH_NS}}}oMath')
    for child in root:
        omath.append(deepcopy(child))
    return omath


def set_paragraph_rtl(paragraph):
    """Add <w:bidi/> to paragraph properties for RTL direction.

    Args:
        paragraph: python-docx Paragraph object.
    """
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def set_run_rtl(run):
    """Add <w:rtl/> to run properties for RTL text.

    Args:
        run: python-docx Run object.
    """
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def _set_run_font(run, font_name=DEFAULT_FONT, font_size=DEFAULT_FONT_SIZE):
    """Set font name and size on a run."""
    run.font.name = font_name
    run.font.size = font_size
    # Also set complex script font for Hebrew
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), font_name)


def add_mixed_paragraph(doc, text, is_rtl=True, style=None):
    """Add paragraph with mixed text and inline LaTeX math.

    Splits text on $...$ delimiters. Even-indexed parts are text runs,
    odd-indexed parts are math (converted to OMML). Falls back to
    italic text on OMML conversion failure.

    Args:
        doc: python-docx Document.
        text: Text string with optional $...$ LaTeX delimiters.
        is_rtl: Whether to set RTL direction.
        style: Optional paragraph style name.

    Returns:
        The created Paragraph object.
    """
    para = doc.add_paragraph(style=style)
    if is_rtl:
        set_paragraph_rtl(para)

    # Split on inline math
    parts = INLINE_MATH_RE.split(text)

    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 0:
            # Text run
            run = para.add_run(part)
            _set_run_font(run)
            if is_rtl:
                set_run_rtl(run)
        else:
            # Math expression
            try:
                omml_elem = latex_to_omml(part)
                para._element.append(omml_elem)
            except Exception:
                # Fallback: italic text
                run = para.add_run(part)
                run.italic = True
                _set_run_font(run)

    return para


def add_plot_image(doc, base64_data, caption, fig_num, language='he'):
    """Embed a base64-encoded plot image in the document.

    Args:
        doc: python-docx Document.
        base64_data: Base64 string, optionally with data URI prefix.
        caption: Caption text.
        fig_num: Figure number.
        language: 'he' or 'en'.

    Raises:
        ValueError: If image data cannot be decoded.
    """
    # Strip data URI prefix
    b64_str = base64_data
    if ',' in base64_data:
        b64_str = base64_data.split(',', 1)[1]

    try:
        img_bytes = base64.b64decode(b64_str)
    except Exception as e:
        raise ValueError(f"Failed to decode plot image: {e}")

    doc.add_picture(BytesIO(img_bytes), width=Inches(5.5))
    # Center the image paragraph
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add caption
    fig_label = 'איור' if language == 'he' else 'Figure'
    caption_para = doc.add_paragraph(f'{fig_label} {fig_num}: {caption}')
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if language == 'he':
        set_paragraph_rtl(caption_para)
    for run in caption_para.runs:
        _set_run_font(run, font_size=Pt(10))


def _add_title_page(doc, title_page, language, is_rtl):
    """Add title page to document.

    Args:
        doc: python-docx Document.
        title_page: Dict with student info and experiment title.
        language: 'he' or 'en'.
        is_rtl: Whether to apply RTL direction.
    """
    exp_title = title_page.get('experimentTitle', '')

    # Title
    title_para = doc.add_paragraph(exp_title, style='Title')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_rtl:
        set_paragraph_rtl(title_para)

    # Labels
    labels = {
        'he': {'name': 'שם', 'id': 'ת.ז.', 'partner': 'שותף/ה למעבדה',
                'partner_id': 'ת.ז. שותף/ה', 'course': 'קורס', 'date': 'תאריך'},
        'en': {'name': 'Name', 'id': 'ID', 'partner': 'Lab Partner',
                'partner_id': 'Partner ID', 'course': 'Course', 'date': 'Date'},
    }
    lbl = labels.get(language, labels['he'])

    fields = [
        (lbl['name'], title_page.get('studentName', '')),
        (lbl['id'], title_page.get('studentId', '')),
        (lbl['partner'], title_page.get('labPartner', '')),
        (lbl['partner_id'], title_page.get('labPartnerId', '')),
        (lbl['course'], title_page.get('courseName', '')),
        (lbl['date'], title_page.get('date', '')),
    ]

    for label_text, value in fields:
        if value:
            para = doc.add_paragraph()
            if is_rtl:
                set_paragraph_rtl(para)
            run = para.add_run(f'{label_text}: {value}')
            _set_run_font(run)
            if is_rtl:
                set_run_rtl(run)

    # Page break
    doc.add_page_break()


def _add_section(doc, heading_text, content, is_rtl):
    """Add a section with heading and mixed content paragraphs.

    Args:
        doc: python-docx Document.
        heading_text: Section heading string.
        content: Section body text (may contain $...$ LaTeX).
        is_rtl: Whether to apply RTL direction.
    """
    heading = doc.add_heading(heading_text, level=2)
    if is_rtl:
        set_paragraph_rtl(heading)

    if not content or not content.strip():
        return

    for para_text in content.split('\n\n'):
        para_text = para_text.strip()
        if para_text:
            # Handle single newlines as line breaks within paragraph
            for line in para_text.split('\n'):
                line = line.strip()
                if line:
                    add_mixed_paragraph(doc, line, is_rtl=is_rtl)


def _add_results_section(doc, analysis_data, plots, language, is_rtl):
    """Add results section with parameter table, stats, and plots.

    Args:
        doc: python-docx Document.
        analysis_data: Dict with fit, formula, nsigma data.
        plots: Dict with fit/residuals base64 images.
        language: 'he' or 'en'.
        is_rtl: Whether to apply RTL direction.
    """
    h = HEADERS.get(language, HEADERS['he'])
    heading = doc.add_heading(h['results'], level=2)
    if is_rtl:
        set_paragraph_rtl(heading)

    fit = analysis_data.get('fit') if analysis_data else None
    if fit and isinstance(fit, dict):
        model = fit.get('modelName', '')
        if model:
            model_lbl = 'מודל התאמה' if language == 'he' else 'Fit Model'
            para = doc.add_paragraph()
            if is_rtl:
                set_paragraph_rtl(para)
            run = para.add_run(f'{model_lbl}: {model}')
            run.bold = True
            _set_run_font(run)

        raw_params = fit.get('parameters', [])
        param_names = fit.get('parameterNames', fit.get('parameter_names', []))
        raw_uncs = fit.get('uncertainties', [])
        raw_rounded = fit.get('rounded', [])

        # Normalize: params can be list of dicts or list of numbers
        if raw_params and isinstance(raw_params[0], dict):
            params = raw_params
        elif raw_params and param_names:
            params = []
            for i, name in enumerate(param_names):
                val = raw_params[i] if i < len(raw_params) else 0
                unc = raw_uncs[i] if i < len(raw_uncs) else 0
                rd = raw_rounded[i] if i < len(raw_rounded) else f'{val} +/- {unc}'
                params.append({'name': name, 'value': val, 'uncertainty': unc, 'rounded': rd})
        else:
            params = []

        if params:
            param_lbl = 'פרמטר' if language == 'he' else 'Parameter'
            value_lbl = 'ערך' if language == 'he' else 'Value'
            unc_lbl = 'אי-ודאות' if language == 'he' else 'Uncertainty'

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = param_lbl
            hdr[1].text = value_lbl
            hdr[2].text = unc_lbl

            for p in params:
                row = table.add_row().cells
                row[0].text = str(p.get('name', ''))
                rounded = str(p.get('rounded', ''))
                if '+/-' in rounded:
                    parts = rounded.split('+/-')
                    row[1].text = parts[0].strip()
                    row[2].text = parts[1].strip()
                elif '±' in rounded:
                    parts = rounded.split('±')
                    row[1].text = parts[0].strip()
                    row[2].text = parts[1].strip()
                else:
                    row[1].text = str(p.get('value', ''))
                    row[2].text = str(p.get('uncertainty', ''))

        # Goodness of fit stats — check both nested goodnessOfFit and flat keys
        gof = fit.get('goodnessOfFit', {}) or {}
        chi_r = gof.get('chiSquaredReduced') or fit.get('chiSquared')
        r_sq = gof.get('rSquared') or fit.get('rSquared')
        p_val = gof.get('pValue') or fit.get('pValue')
        dof = gof.get('dof') or fit.get('degreesOfFreedom')
        gof_parts = []
        if chi_r is not None and dof is not None:
            gof_parts.append(f"chi^2/dof = {float(chi_r) / float(dof):.4f}")
        elif chi_r is not None:
            gof_parts.append(f"chi^2 = {float(chi_r):.4f}")
        if r_sq is not None:
            gof_parts.append(f"R^2 = {float(r_sq):.6f}")
        if p_val is not None:
            gof_parts.append(f"P = {float(p_val):.4f}")
        if gof_parts:
            para = doc.add_paragraph(', '.join(gof_parts))
            if is_rtl:
                set_paragraph_rtl(para)

    # Formula
    formula = analysis_data.get('formula') if analysis_data else None
    if formula and isinstance(formula, dict):
        expr = formula.get('expression', '')
        formatted = formula.get('formatted', '')
        if expr and formatted:
            formula_lbl = 'חישוב נוסחה' if language == 'he' else 'Formula Calculation'
            add_mixed_paragraph(doc, f'{formula_lbl}: ${expr} = {formatted}$', is_rtl=is_rtl)

    # N-sigma
    nsigma = analysis_data.get('nsigma') if analysis_data else None
    if nsigma and isinstance(nsigma, dict):
        ns = nsigma.get('nSigma', '')
        verdict = nsigma.get('verdict', '')
        if ns != '' and verdict:
            nsigma_lbl = 'השוואת N-sigma' if language == 'he' else 'N-sigma Comparison'
            add_mixed_paragraph(doc, f'{nsigma_lbl}: $N_\\sigma = {ns}$ - {verdict}', is_rtl=is_rtl)

    # Plot images
    fig_num = 1
    for plot_key, caption_he, caption_en in [
        ('fit', 'גרף התאמה', 'Fit Plot'),
        ('residuals', 'שאריות', 'Residuals'),
    ]:
        img_data = plots.get(plot_key) if plots else None
        if img_data and isinstance(img_data, str):
            caption = caption_he if language == 'he' else caption_en
            add_plot_image(doc, img_data, caption, fig_num, language=language)
            fig_num += 1


def _setup_document(language='he'):
    """Create and configure a new Document with fonts and margins.

    Returns:
        Tuple of (Document, is_rtl boolean).
    """
    doc = Document()
    is_rtl = language == 'he'

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = DEFAULT_FONT
    font.size = DEFAULT_FONT_SIZE
    # Set complex script font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:cs'), DEFAULT_FONT)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    return doc, is_rtl


def generate_docx(sections, title_page, plots, analysis_data, language='he'):
    """Generate a full lab report DOCX.

    Args:
        sections: Dict with keys theory, method, discussion, conclusions.
        title_page: Dict with student info and experiment title.
        plots: Dict with fit/residuals base64 images.
        analysis_data: Dict with fit, formula, nsigma data.
        language: 'he' or 'en'.

    Returns:
        DOCX file contents as bytes.
    """
    doc, is_rtl = _setup_document(language)
    h = HEADERS.get(language, HEADERS['he'])

    # Title page
    _add_title_page(doc, title_page, language, is_rtl)

    # Theory
    _add_section(doc, h['theory'], sections.get('theory', ''), is_rtl)

    # Method
    _add_section(doc, h['method'], sections.get('method', ''), is_rtl)

    # Results
    _add_results_section(doc, analysis_data, plots, language, is_rtl)

    # Discussion
    _add_section(doc, h['discussion'], sections.get('discussion', ''), is_rtl)

    # Conclusions
    _add_section(doc, h['conclusions'], sections.get('conclusions', ''), is_rtl)

    # Save to bytes
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_results_docx(analysis_data, plots, summary='', language='he'):
    """Generate a results-only DOCX (no title page, no AI sections).

    Args:
        analysis_data: Dict with fit, formula, nsigma data.
        plots: Dict with fit/residuals base64 images.
        summary: AI-generated summary text.
        language: 'he' or 'en'.

    Returns:
        DOCX file contents as bytes.
    """
    doc, is_rtl = _setup_document(language)

    # Heading
    heading_text = 'תוצאות ניתוח' if language == 'he' else 'Analysis Results'
    title = doc.add_heading(heading_text, level=1)
    if is_rtl:
        set_paragraph_rtl(title)

    # Summary
    if summary and summary.strip():
        add_mixed_paragraph(doc, summary.strip(), is_rtl=is_rtl)

    # Results section content (without heading - we already have the title)
    fit = analysis_data.get('fit') if analysis_data else None
    if fit and isinstance(fit, dict):
        model = fit.get('modelName', '')
        if model:
            model_lbl = 'מודל התאמה' if language == 'he' else 'Fit Model'
            para = doc.add_paragraph()
            if is_rtl:
                set_paragraph_rtl(para)
            run = para.add_run(f'{model_lbl}: {model}')
            run.bold = True
            _set_run_font(run)

        params = fit.get('parameters', [])
        if params:
            param_lbl = 'פרמטר' if language == 'he' else 'Parameter'
            value_lbl = 'ערך' if language == 'he' else 'Value'
            unc_lbl = 'אי-ודאות' if language == 'he' else 'Uncertainty'

            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = param_lbl
            hdr[1].text = value_lbl
            hdr[2].text = unc_lbl

            for p in params:
                row = table.add_row().cells
                row[0].text = p.get('name', '')
                rounded = p.get('rounded', '')
                if rounded:
                    parts = rounded.split('+/-')
                    if len(parts) == 2:
                        row[1].text = parts[0].strip()
                        row[2].text = parts[1].strip()
                    else:
                        row[1].text = rounded
                        row[2].text = ''
                else:
                    row[1].text = str(p.get('value', ''))
                    row[2].text = str(p.get('uncertainty', ''))

        # Goodness of fit
        gof = fit.get('goodnessOfFit', {})
        gof_parts = []
        if gof.get('chiSquaredReduced') is not None:
            gof_parts.append(f"chi^2/dof = {gof['chiSquaredReduced']:.4f}")
        if gof.get('rSquared') is not None:
            gof_parts.append(f"R^2 = {gof['rSquared']:.6f}")
        if gof.get('pValue') is not None:
            gof_parts.append(f"P = {gof['pValue']:.4f}")
        if gof_parts:
            para = doc.add_paragraph(', '.join(gof_parts))
            if is_rtl:
                set_paragraph_rtl(para)

    # N-sigma
    nsigma = analysis_data.get('nsigma') if analysis_data else None
    if nsigma and isinstance(nsigma, dict):
        ns = nsigma.get('nSigma', '')
        verdict = nsigma.get('verdict', '')
        if ns != '' and verdict:
            nsigma_lbl = 'השוואת N-sigma' if language == 'he' else 'N-sigma Comparison'
            add_mixed_paragraph(doc, f'{nsigma_lbl}: $N_\\sigma = {ns}$ - {verdict}', is_rtl=is_rtl)

    # Plots
    fig_num = 1
    for plot_key, caption_he, caption_en in [
        ('fit', 'גרף התאמה', 'Fit Plot'),
        ('residuals', 'שאריות', 'Residuals'),
    ]:
        img_data = plots.get(plot_key) if plots else None
        if img_data and isinstance(img_data, str):
            caption = caption_he if language == 'he' else caption_en
            add_plot_image(doc, img_data, caption, fig_num, language=language)
            fig_num += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
