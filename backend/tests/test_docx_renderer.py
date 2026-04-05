"""Tests for the DOCX rendering pipeline."""
import pytest
from lxml import etree

# Minimal 1x1 PNG as base64 for plot testing
TINY_PNG_B64 = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="

SAMPLE_SECTIONS = {
    "theory": "According to Hooke's law, $F = -kx$, where $k$ is the spring constant.\n\nThe restoring force is proportional to displacement.",
    "method": "We measured force vs displacement using a spring and force sensor.\n\nMasses from 50g to 500g were applied.",
    "discussion": "The measured value $k = 49.8 \\pm 0.5 \\text{ N/m}$ agrees with the expected value.\n\nError sources include friction.",
    "conclusions": "The spring constant was measured as $k = 49.8 \\pm 0.5 \\text{ N/m}$, consistent with theory.",
}

SAMPLE_TITLE_PAGE = {
    "studentName": "Test Student",
    "studentId": "123456789",
    "labPartner": "Partner Name",
    "courseName": "Physics Lab 1",
    "experimentTitle": "Hooke's Law",
    "date": "2026-03-27",
}

SAMPLE_ANALYSIS_DATA = {
    "fit": {
        "modelName": "linear",
        "parameters": [
            {"name": "k", "value": 49.8, "uncertainty": 0.5, "rounded": "49.8 +/- 0.5", "latex": "k"},
            {"name": "b", "value": 0.12, "uncertainty": 0.03, "rounded": "0.12 +/- 0.03", "latex": "b"},
        ],
        "goodnessOfFit": {
            "chiSquaredReduced": 1.23,
            "rSquared": 0.9987,
            "pValue": 0.45,
            "dof": 8,
        },
    },
}

MATH_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/math}'
WORD_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class TestLatexToOmml:
    """Tests for latex_to_omml conversion."""

    def test_latex_to_omml(self):
        """Basic LaTeX -> OMML returns oMath element."""
        from utils.docx_renderer import latex_to_omml
        elem = latex_to_omml("E = mc^2")
        assert elem is not None
        assert elem.tag == f'{MATH_NS}oMath'

    def test_latex_to_omml_fraction(self):
        """Fraction LaTeX produces m:f sub-element."""
        from utils.docx_renderer import latex_to_omml
        elem = latex_to_omml(r"\frac{a}{b}")
        xml_str = etree.tostring(elem, encoding='unicode')
        assert f'{MATH_NS}f' in xml_str or 'm:f' in xml_str

    def test_latex_to_omml_fallback(self):
        """Empty LaTeX raises exception."""
        from utils.docx_renderer import latex_to_omml
        with pytest.raises(Exception):
            latex_to_omml("")


class TestRtlBidi:
    """Tests for RTL bidi support."""

    def test_rtl_bidi(self):
        """set_paragraph_rtl adds w:bidi element."""
        from docx import Document
        from utils.docx_renderer import set_paragraph_rtl
        doc = Document()
        para = doc.add_paragraph("test")
        set_paragraph_rtl(para)
        xml_str = etree.tostring(para._element, encoding='unicode')
        assert 'w:bidi' in xml_str


class TestPlotEmbedding:
    """Tests for plot image embedding."""

    def test_plot_embedding(self):
        """add_plot_image results in inline image relationship."""
        from docx import Document
        from utils.docx_renderer import add_plot_image
        doc = Document()
        add_plot_image(doc, TINY_PNG_B64, "Fit Plot", 1, language='en')
        # Check that there's at least one inline shape (image)
        rels = doc.part.rels
        image_rels = [r for r in rels.values() if 'image' in r.reltype]
        assert len(image_rels) > 0


class TestFullReportDocx:
    """Tests for full report generation."""

    def test_full_report_docx(self):
        """generate_docx returns valid DOCX with Hebrew headers."""
        from utils.docx_renderer import generate_docx
        from docx import Document
        from io import BytesIO

        docx_bytes = generate_docx(
            sections=SAMPLE_SECTIONS,
            title_page=SAMPLE_TITLE_PAGE,
            plots={"fit": TINY_PNG_B64},
            analysis_data=SAMPLE_ANALYSIS_DATA,
            language='he',
        )
        # Valid zip (DOCX is a zip)
        assert docx_bytes[:2] == b'PK'

        # Parse and check Hebrew headers present
        doc = Document(BytesIO(docx_bytes))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'רקע תיאורטי' in all_text
        assert 'שיטת מדידה' in all_text
        assert 'תוצאות' in all_text


class TestResultsOnlyDocx:
    """Tests for results-only report generation."""

    def test_results_only_docx(self):
        """generate_results_docx returns valid DOCX with results."""
        from utils.docx_renderer import generate_results_docx
        from docx import Document
        from io import BytesIO

        docx_bytes = generate_results_docx(
            analysis_data=SAMPLE_ANALYSIS_DATA,
            plots={"fit": TINY_PNG_B64},
            summary="This is a test summary.",
            language='he',
        )
        assert docx_bytes[:2] == b'PK'

        doc = Document(BytesIO(docx_bytes))
        all_text = '\n'.join(p.text for p in doc.paragraphs)
        assert 'test summary' in all_text.lower() or 'This is a test summary' in all_text


class TestMixedTextMath:
    """Tests for mixed text and math paragraphs."""

    def test_mixed_text_math(self):
        """add_mixed_paragraph creates paragraph with text and OMML."""
        from docx import Document
        from utils.docx_renderer import add_mixed_paragraph

        doc = Document()
        para = add_mixed_paragraph(doc, "the value is $k = 50$ N/m")
        xml_str = etree.tostring(para._element, encoding='unicode')
        # Should contain both text runs and OMML math
        assert 'the value is' in xml_str or 'w:t' in xml_str
        assert f'{MATH_NS}oMath' in xml_str or 'm:oMath' in xml_str
